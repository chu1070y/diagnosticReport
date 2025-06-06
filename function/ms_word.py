import re
import pandas as pd

from function.db_work import DBwork
from function.parser import Parser
from module.common import Common

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def find_table_after_heading(doc, heading_text):
    """ 특정 제목(heading_text) 다음에 오는 첫 번째 표 찾기 """
    found_heading = False
    for para in doc.paragraphs:
        if heading_text in para.text.strip():  # 제목 찾기
            found_heading = True
            continue

        if found_heading:
            # 제목 다음에 있는 표 반환
            for table in doc.tables:
                if para._element in table._element.xpath(".//preceding-sibling::w:p"):
                    return table
    return None


def set_cell_font(cell, bold=False, center=False, font_size=8):
    """ 셀의 폰트를 설정하는 함수 """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(font_size)
            run.font.bold = bold
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if center else WD_PARAGRAPH_ALIGNMENT.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, color):
    """ 셀의 배경색을 설정하는 함수 """
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")  # 값 설정 (필수)
    shading.set(qn("w:color"), "auto")  # 텍스트 색상 자동
    shading.set(qn("w:fill"), color)  # 배경색 설정

    # 셀에 적용
    tc_pr = cell._element.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell._element.append(tc_pr)
    tc_pr.append(shading)


def parse_table_to_dict(table_str):
    lines = table_str.strip().split("\n")  # 문자열을 줄 단위로 나눔
    data_dict = {}

    for line in lines:
        line = line.strip()

        # 구분선 제거
        if line.startswith("+") or line.startswith("| Variable_name"):
            continue

        # 정규 표현식을 사용하여 컬럼 데이터 추출
        match = re.match(r"\|\s*(.*?)\s*\|\s*(.*?)\s*\|", line)
        if match:
            key, value = match.groups()
            data_dict[key.strip()] = value.strip()

    return data_dict


def parse_cpu_info(cpu_text):
    from collections import defaultdict

    model_name = None
    physical_cores = defaultdict(int)  # 물리 CPU별 물리 코어 수
    logical_cores = defaultdict(int)   # 물리 CPU별 논리 코어 수

    # 프로세서 블록 단위로 나누기
    cpu_blocks = cpu_text.strip().split("\n\n")

    for block in cpu_blocks:
        lines = block.strip().split("\n")
        cpu_info = {}

        for line in lines:
            match = re.match(r"([^:]+)\s*:\s*(.+)", line)
            if match:
                key, value = match.groups()
                cpu_info[key.strip()] = value.strip()

        # CPU 모델명 추출 (모든 프로세서가 동일한 모델 가정)
        if "model name" in cpu_info and model_name is None:
            model_name = cpu_info["model name"]

        # 물리 ID 기준으로 물리/논리 코어 개수 저장
        if "physical id" in cpu_info:
            phys_id = cpu_info["physical id"]
            if "cpu cores" in cpu_info:
                physical_cores[phys_id] = int(cpu_info["cpu cores"])
            if "siblings" in cpu_info:
                logical_cores[phys_id] = int(cpu_info["siblings"])

    # 총 물리 코어 & 논리 코어 개수 계산
    total_physical_cores = sum(physical_cores.values())
    total_logical_cores = sum(logical_cores.values())

    return f"{model_name} ({total_physical_cores} Physical Cores / {total_logical_cores} Logical Cores)"


def make_table_variables_info(report, paragraph, variables_data):
    table = report.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    paragraph._element.addnext(table._element)

    # 헤더
    table.cell(0, 0).text = ""
    table.cell(0, 0).paragraphs[0].add_run('Variable_name')
    table.cell(0, 1).text = ""
    table.cell(0, 1).paragraphs[0].add_run('Value')

    set_cell_shading(table.cell(0, 0), "D9E1F2")
    set_cell_shading(table.cell(0, 1), "D9E1F2")

    set_cell_font(table.cell(0, 0), bold=True, center=True)
    set_cell_font(table.cell(0, 1), bold=True, center=True)

    # 2행부터 데이터 추가
    for i, (key, value) in enumerate(variables_data.items()):
        row_cells = table.add_row().cells  # 새로운 행 추가

        row_cells[0].paragraphs[0].add_run(key)
        row_cells[1].paragraphs[0].add_run(str(value))

        # 줄무늬 스타일 (홀수 행은 회색, 짝수 행은 흰색)
        set_cell_shading(row_cells[0], "F2F2F2")

        # 폰트 스타일
        set_cell_font(row_cells[0])
        set_cell_font(row_cells[1])


def make_table_count(report, paragraph, data: list, header=[], use_column=[]):
    if use_column:
        data = [
            tuple(val for val, use in zip(row, use_column) if use == 1)
            for row in data
        ]

    header = header if header else data[0]
    data = data[1:]

    df = pd.DataFrame(data, columns=header)
    grouped = df.groupby(header[0])

    current = paragraph._element

    for schema, group in grouped:
        table = report.add_table(rows=1, cols=len(header))
        table.style = 'Table Grid'
        table.autofit = False
        current.addnext(table._element)

        # 헤더
        for i in range(len(header)):
            cell = table.rows[0].cells[i]

            cell.text = header[i]
            set_cell_shading(cell, "D9E1F2")
            set_cell_font(cell, bold=True, center=True, font_size=8)

        # 데이터
        for _, row in group.iterrows():
            cells = table.add_row().cells

            for i in range(len(header)):
                cells[i].text = str(row[header[i]])
                set_cell_font(cells[i], center=True, font_size=8)

        # 첫 번째 컬럼 병합
        row_count = len(group)
        if row_count > 1:
            start_cell = table.cell(1, 0)
            end_cell = table.cell(row_count, 0)
            start_cell.merge(end_cell)

            table.cell(1, 0).text = schema
            set_cell_font(table.cell(1, 0), center=True, font_size=8)

        next_para = report.add_paragraph()
        next_para.add_run("")

        current = table._element
        current.addnext(next_para._element)
        current = next_para._element


class MSword(Common):
    def __init__(self):
        self.logger = self.get_logger()
        self.report_conf = self.get_config()['excel']

        self.sample_report_path = './sample/sample_diagnostic_report_v1.3.docx'
        self.report_path = './result/diagnostic_report.docx'
        self.os_info_path = self.get_config()['path'].get('os_info_file')
        self.db_info_path = self.get_config()['path'].get('db_info_file')

        self.graph_folder = './result/graphs/'
        # self.image_extensions = [".png", ".jpg", ".jpeg"]

        self.parser = Parser()

    @Common.exception_handler
    def make_report_v2(self):
        # ################### test
        # self.sample_report_path = '../sample/sample_diagnostic_report_v1.3.docx'
        # self.report_path = './result/diagnostic_report.docx'
        # self.os_info_path = self.get_config()['path'].get('os_info_file')
        # self.graph_folder = './result/graphs/'
        # ###############################

        # 설정값 + 상태값 + 그래프 삽입
        self.logger.info("input status and graph on diagnostic report")
        db = DBwork()
        status_data = db.get_latest_status_data()

        # 왼쪽 정렬 파라미터 리스트 (여기 없으면 모두 가운데 정렬)
        left_aligned_param_list = [
            'my.cnf', 'hostname', 'memory', 'os version', 'cpu', 'version',
            'log_error', 'slow_query_log_file', 'log_bin_basename', 'innodb_data_file_path',
            'default_storage_engine', 'port', 'socket', 'basedir', 'datadir'
        ]

        ############## OS info
        os_info_data = dict()

        variables_data = dict()

        try:
            os_info_dict = self.parser.parse_osinfo(self.os_info_path)
            variables_data = self.parser.parse_table_to_dict(os_info_dict.get('global variables'))

            os_info_data['hostname'] = os_info_dict.get('hostname')
            os_info_data['OS Version'] = os_info_dict.get('OS Version').split('\n')[0]
            os_info_data['CPU'] = parse_cpu_info(os_info_dict.get('cpu'))
            os_info_data['Memory'] = os_info_dict.get('memory').split('\n')[0].replace('MemTotal:', '').strip()
            os_info_data['my.cnf'] = os_info_dict.get('my.cnf')

        except Exception as e:
            self.logger.warning("An error occurred while parsing the OS info file...")
            self.logger.warning(e)

        ############## DB info
        db_info_dict = dict()

        schema_table_data = []
        engine_table_data = []
        unused_user_data = []
        unused_index_data = []
        duplicate_index_data = []
        low_cardinality_index_data = []
        no_update_tables_data = []
        fullscan_sql_data = []
        temptable_sql_data = []
        unconstrained_tables_data = []

        try:
            db_info_dict = self.parser.parse_dbinfo(self.db_info_path)

            schema_table_data = self.parser.parse_table_to_list(db_info_dict.get('3.5.5.2 스키마별 테이블 사이즈'))
            engine_table_data = self.parser.parse_table_to_list(db_info_dict.get('3.5.5.3 스토리지 엔진별 테이블 개수'))
            unused_user_data = self.parser.parse_table_to_list(db_info_dict.get('5.1 미사용DB계정 확인'))
            unused_index_data = self.parser.parse_table_to_list(db_info_dict.get('5.2.1 미사용 인덱스'))
            duplicate_index_data = self.parser.parse_table_to_list(db_info_dict.get('5.2.2 중복 인덱스'))
            low_cardinality_index_data = self.parser.parse_table_to_list(db_info_dict.get('5.2.3 분포도 낮은 인덱스'))
            no_update_tables_data = self.parser.parse_table_to_list(db_info_dict.get('5.3 변경 없는 테이블 진단'))
            fullscan_sql_data = self.parser.parse_table_to_list(db_info_dict.get('5.4.1 Full Table Scan SQL'))
            temptable_sql_data = self.parser.parse_table_to_list(db_info_dict.get('5.4.2 임시테이블 사용 SQL'))
            unconstrained_tables_data = self.parser.parse_table_to_list(db_info_dict.get('5.5 PK/UK 누락 테이블'))

        except Exception as e:
            self.logger.warning("An error occurred while parsing the DB info file...")
            self.logger.warning(e)

        all_data = {**status_data, **os_info_data, **variables_data}
        all_data = {k.lower(): v for k, v in all_data.items()}

        report = Document(self.sample_report_path)

        ################## 표 내부 데이터 변경
        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()

                    if text.startswith("{graph_") and text.endswith("}"):
                        self.insert_graph(self.graph_folder, cell.paragraphs[0], text)

                    elif text.startswith("{"):
                        key = text.lstrip('{').rstrip('}').lower()

                        if key in all_data.keys():
                            cell.text = ""  # 기존 텍스트 삭제 (안 하면 스타일 초기화됨)
                            para = cell.paragraphs[0]

                            run = para.add_run(all_data.get(key))
                            run.font.name = "Malgun Gothic"

                            # 정렬 방식 지정
                            if key in left_aligned_param_list:
                                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                            else:
                                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ################## 표 외부 데이터 변경
        for paragraph in report.paragraphs:
            words = paragraph.text.split()
            for word in words:
                # 이미지 삽입 처리
                if word.startswith("{graph_") and word.endswith("}"):
                    self.insert_graph(self.graph_folder, paragraph, word)

                elif word.startswith("{_table_") and word.endswith("}"):
                    paragraph.clear()

                    ### [표 생성] 로직
                    if word == '{_table_variables_info}':
                        make_table_variables_info(report, paragraph, variables_data)

                    elif word == '{_table_schema_table}':
                        make_table_count(report, paragraph, schema_table_data, use_column=[1,1,1], header=['스키마 이름', '테이블 이름', '크기 (GB)'])

                    elif word == '{_table_engine_table}':
                        make_table_count(report, paragraph, engine_table_data, use_column=[1,1,1], header=['스키마 이름', '스토리지 엔진', '개수'])

                    elif word == '{_table_unused_user}':
                        self.make_table_auto(report, paragraph, unused_user_data, width=[0.2, 0.2])

                    elif word == '{_table_unused_index}':
                        self.make_table_auto(report, paragraph, unused_index_data, header=['스키마 이름', '테이블 이름', '인덱스 이름'])

                    elif word == '{_table_duplicate_index}':
                        self.make_table_auto(report, paragraph, duplicate_index_data, header=['스키마 이름', '테이블 이름', 'redundant index', 'redundant index column', 'dominanat index', 'dominant index column'])

                    elif word == '{_table_low_cardinality_index}':
                        self.make_table_auto(report, paragraph, low_cardinality_index_data, use_column=[1,1,1,1,0,0,0,0,1], header=['스키마 이름', '테이블 이름', '인덱스 이름', '컬럼 이름', '분포도율(%)'])

                    elif word == '{_table_no_update_tables}':
                        self.make_table_auto(report, paragraph, no_update_tables_data, header=['스키마 이름', '테이블 이름', '레코드 수', 'Count_read', 'Count_write'])

                    elif word == '{_table_fullscan_sql}':
                        self.make_table_auto(report, paragraph, fullscan_sql_data, use_column=[1,1,0,0,0,0,0], header=['스키마 이름', 'SQL'], width=[0.2, 0.8])

                    elif word == '{_table_temptable_sql}':
                        self.make_table_auto(report, paragraph, temptable_sql_data, use_column=[1,1,0,0,0], header=['스키마 이름', 'SQL'], width=[0.2, 0.8])

                    elif word == '{_table_unconstrained_tables}':
                        self.make_table_auto(report, paragraph, unconstrained_tables_data, header=['스키마 이름', '테이블 이름'], width=[0.15, 0.4])

        # 수정된 문서 저장
        report.save(self.report_path)

    def insert_graph(self, graph_folder, paragraph, placeholder):
        """
        문단 내 플레이스홀더({graph_파일명})를 찾아 이미지를 삽입하는 함수
        """
        filename = placeholder.replace('{graph_', '').strip('}')
        image_path = None
        image_extensions = [".png"]

        # 해당 파일명이 존재하는 이미지 찾기
        for ext in image_extensions:
            potential_path = f"{graph_folder}{filename}{ext}"
            try:
                with open(potential_path, "rb"):  # 파일 존재 확인
                    image_path = potential_path
                    break
            except FileNotFoundError:
                continue

        if image_path:
            paragraph.text = ""  # 기존 텍스트 제거
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))  # 이미지 크기 조정
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 가운데 정렬 적용
        else:
            self.logger.warning(f"이미지 파일을 찾을 수 없음: {filename}")

    @Common.exception_handler
    def get_os_info(self):
        with open(self.os_info_path, 'r', encoding='utf-8') as file:
            content = file.read()

        sections = re.split(r"=+\s*(.+?)\s*=+\n", content)
        session_data = {sections[i].strip(): sections[i + 1].strip() for i in range(1, len(sections) - 1, 2)}

        return session_data

    @staticmethod
    @Common.exception_handler
    def make_table_auto(report, paragraph, data: list, use_column=[], header=[], width=[]):
        if use_column:
            data = [
                tuple(val for val, use in zip(row, use_column) if use == 1)
                for row in data
            ]

        table = report.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.autofit = False
        paragraph._element.addnext(table._element)

        total_width = 6.25
        col_widths = [total_width * r for r in width]

        for row_idx, row_data in enumerate(data):
            for col_idx, value in enumerate(row_data):

                cell = table.cell(row_idx, col_idx)
                cell.text = str(value)

                if width:
                    cell.width = Inches(col_widths[col_idx])

                if row_idx == 0:
                    if header:
                        cell.text = header[col_idx]

                    set_cell_shading(cell, "D9E1F2")
                    set_cell_font(cell, bold=True, center=True, font_size=8)

                else:
                    set_cell_font(cell, font_size=8)


if __name__ == "__main__":
    # MSword().make_report()
    MSword().make_report_v2()
